import os
import pymupdf
import pytesseract
from PIL import Image
import io
from docx import Document


# Если Tesseract не в PATH, укажите путь:
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def convert_file_to_text(file_path, use_ocr=True):
    if not os.path.exists(file_path):
        return "Ошибка: Файл не найден"

    ext = file_path.lower().split('.')[-1]

    try:
        # --- PDF (с поддержкой OCR) ---
        if ext == 'pdf':
            text = ""
            with pymupdf.open(file_path) as doc:
                for page in doc:
                    page_text = page.get_text().strip()
                    if not page_text and use_ocr:
                        pix = page.get_pixmap(matrix=pymupdf.Matrix(2, 2))
                        img_data = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_data))
                        page_text = pytesseract.image_to_string(img, lang='rus+eng')
                    text += page_text + "\n"
            return text

        # --- DOCX (ИСПРАВЛЕНО: ТЕПЕРЬ ЧИТАЕТ ТАБЛИЦЫ) ---
        elif ext == 'docx':
            doc = Document(file_path)
            full_text = []

            # 1. Извлекаем текст из обычных абзацев
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # 2. Извлекаем текст из всех ячеек всех таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Убираем дублирование текста, если ячейки объединены
                        cell_text = cell.text.strip()
                        if cell_text and (not full_text or cell_text != full_text[-1]):
                            full_text.append(cell_text)

            return "\n".join(full_text)

        # --- TXT ---
        elif ext == 'txt':
            for encoding in ['utf-8', 'cp1251']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            return "Ошибка: Кодировка не поддерживается"

        else:
            return f"Ошибка: Формат .{ext} не поддерживается"

    except Exception as e:
        return f"Ошибка при обработке файла: {str(e)}"